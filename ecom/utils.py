import datetime
from django.utils import timezone
from dateutil.relativedelta import relativedelta
import json
from django.http import JsonResponse
import arrow
from CRMS.settings import REDIS
from ecom.models import Product
from django.http import JsonResponse
import arrow
from ecom.models import Product, Purchase, History
from django.db import transaction
from CRMS.notif import Notify
from django.db import transaction
from django.db.models import Sum, Avg, Count
from django.db.models.functions import TruncMonth, TruncYear
import logging
from docx import Document
from docx.shared import Pt
from docx.shared import Inches

logger = logging.getLogger(__name__)

notify = Notify()

"""

TODO:Method to generate a report of low stock products  
     Method to generate sales analytics
     Such as the best-selling products, top revenue-generating products, 
"""


class Inventory:
    def get_products(self):
        try:
            products = Product.objects.values(
                "product_uid",
                "product_name",
                "price",
                "quantity",
                "description",
                "category",
                "expiry_date",
                "manufacturer",
                "supplier_info",
            ).all()
            if products:
                return JsonResponse({"success": True, "data": list(products)})
            else:
                return JsonResponse({"success": False, "info": "No products found"})
        except Exception as e:
            logger.warning(str(e))
            return JsonResponse(
                {"success": False, "info": "Kindly try again --p2prx2--"}
            )

    def add_product(
        self,
        product_name,
        price,
        quantity,
        expiry_date,
        manufacturer,
        supplier_info,
        category,
        description,
    ):
        try:
            if Product.objects.filter(product_name=product_name).exists():
                return JsonResponse({"error": "Product already exists"}, status=409)

            # Create a new product object
            product = Product.objects.create(
                product_name=product_name,
                price=price,
                quantity=quantity,
                expiry_date=expiry_date,
                manufacturer=manufacturer,
                supplier_info=supplier_info,
                category=category,
                description=description,
                added_on=arrow.now().datetime,
            )

            if product:
                return JsonResponse(
                    {"success": "Product added successfully"}, status=200
                )
            else:
                return JsonResponse({"error": "Failed to add product"}, status=500)

        except Exception as e:
            logger.warning(str(e))
            return JsonResponse(
                {"success": False, "info": "Kindly try again --p2prx2--"}
            )

    @transaction.atomic()
    def update_product(
        self,
        product_uid,
        product_name=None,
        price=None,
        expiry=None,
        quantity=None,
        supplier_info=None,
        description=None,
        category=None,
    ):
        try:
            product = Product.objects.get(product_uid=product_uid)

            if product_name is not None:
                product.product_name = product_name
            if price is not None:
                product.price = price
            if expiry is not None:
                product.expiry_date = expiry
            if quantity is not None:
                product.quantity = quantity
            if supplier_info is not None:
                product.supplier_info = supplier_info
            if description is not None:
                product.description = description
            if category is not None:
                product.category = category

            product.save()

            return JsonResponse({"success": "Product updated successfully"}, status=200)
        except Exception as e:
            logger.warning(str(e))
            return JsonResponse(
                {"success": False, "info": " Unable to update product"}, status=500
            )

    def delete_product(self, product_uid):
        try:
            product = Product.objects.get(product_uid=product_uid)
            product.delete()
            return JsonResponse({"success": "Product deleted successfully"}, status=200)
        except Exception as e:
            logger.warning(str(e))
            return JsonResponse(
                {"success": False, "info": "Unable to delete product"}, status=500
            )

    def search_product(self, query):
        try:
            products = Product.objects.values(
                "product_uid",
                "product_name",
                "price",
                "quantity",
                "description",
                "category",
                "expiry_date",
                "manufacturer",
                "supplier_info",
                "added_on",
            ).filter(product_name__icontains=query)
            if products:
                return JsonResponse({"success": True, "data": list(products)})
            else:
                return JsonResponse({"success": False, "info": "No products found"})
        except Exception as e:
            logger.warning(str(e))
            return JsonResponse(
                {"success": False, "info": "Kindly try again --p2prx2--"}
            )

    def filter_products(self, category=None, manufacturer=None):
        try:
            filter_params = {}
            if category:
                filter_params["category"] = category
            if manufacturer:
                filter_params["manufacturer"] = manufacturer

            products = Product.objects.values(
                "product_uid",
                "product_name",
                "price",
                "quantity",
                "description",
                "category",
                "expiry_date",
                "manufacturer",
                "supplier_info",
                "added_on",
            ).filter(**filter_params)

            if products:
                return JsonResponse({"success": True, "data": list(products)})
            else:
                return JsonResponse({"success": False, "info": "No products found"})

        except Exception as e:
            logger.warning(str(e))
            return JsonResponse(
                {"success": False, "info": "Kindly try again --p2prx2--"}
            )

    def check_expiration(self):
        current_date = datetime.datetime.now().date()
        products = Product.objects.all()
        expired_products = []
        expiring_products = []

        for product in products:
            expiry_date = product.expiry_date.date()
            six_months_from_now = current_date + relativedelta(months=6)

            if current_date > expiry_date:
                json_data = {
                    "product_name": product.product_name,
                    "price": str(product.price),
                    "expiry_date": product.expiry_date.date().isoformat(),
                }
                expired_products.append(json_data)
                REDIS.set("expired-products", json.dumps(expired_products))

            if current_date <= expiry_date <= six_months_from_now:

                time_until_expiration = (expiry_date - current_date).days
                json_data = {
                    "product_name": product.product_name,
                    "price": str(product.price),
                    "expiry_date": product.expiry_date.date().isoformat(),
                    "time_until_expiration": f"{time_until_expiration} days",
                }
                expiring_products.append(json_data)
                REDIS.set("almost-expired-products", json.dumps(expiring_products))

        logger.warning(f"Found {len(expired_products)} expired products.")
        logger.warning(f"Found {len(expiring_products)} expiring soon products.")

    def fetch_total_purchases(self):
        purchases = Purchase.objects.values("product__product_name").annotate(
            total_purchase_amount=Sum("purchase_price"),
            total_quantities_sold=Sum("quantity"),
        )
        if purchases:
            return JsonResponse({"success": True, "data": list(purchases)})
        else:
            return JsonResponse({"success": False, "info": "No purchases found"})

    def print_annual_and_monthly_purchases(self):
        current_month_purchases = (
            Purchase.objects.values("product__product_name")
            .annotate(month=TruncMonth("date"))
            .values("product__product_name", "month")
            .annotate(
                total_purchase_amount=Sum("purchase_price"),
                total_quantities_sold=Sum("quantity"),
            )
            .order_by("product__product_name", "month")
        )

        current_year_purchases = (
            Purchase.objects.values("product__product_name")
            .annotate(year=TruncYear("date"))
            .values("product__product_name", "year")
            .annotate(
                total_purchase_amount=Sum("purchase_price"),
                total_quantities_sold=Sum("quantity"),
            )
            .order_by("product__product_name", "year")
        )

        if current_month_purchases or current_year_purchases:
            return JsonResponse(
                {
                    "success": True,
                    "current_month_purchases": list(current_month_purchases),
                    "current_year_purchases": list(current_year_purchases),
                }
            )
        else:
            return JsonResponse({"success": False, "info": "No purchases found"})

    def register_purchase(self, prod_uid, quantity):
        try:
            product = Product.objects.get(product_uid=prod_uid)

            if product.quantity < int(quantity):
                return JsonResponse({"success": False, "info": "Insufficient stock"})

            if not product.on_sale:
                return JsonResponse(
                    {"success": False, "info": "Product is not on sale"}
                )

            exp_prod = REDIS.get("expired-products")
            expired = json.loads(exp_prod)

            cached_products = [
                {
                    "product_name": item["product_name"],
                    "price": item["price"],
                    "expiry_date": item["expiry_date"],
                }
                for item in expired
            ]

            if product.product_name in [
                item["product_name"] for item in cached_products
            ]:
                return JsonResponse({"success": False, "info": "Product is expired"})

            with transaction.atomic():
                product.quantity -= int(quantity)
                purchase_price = product.price * int(quantity)
                purchase = Purchase.objects.create(
                    product=product, quantity=quantity, purchase_price=purchase_price
                )
                product.save(update_fields=["quantity"])

            if purchase:
                self.generate_invoice(
                    purchase_uid=purchase.purchase_uid, cashier="admin"
                )
                return JsonResponse({"success": True, "info": "Purchase successful"})
            else:
                return JsonResponse({"success": False, "info": "Purchase failed"})

        except Product.DoesNotExist:
            return JsonResponse({"success": False, "info": "Product not found"})
        except Exception as e:
            logger.warning(str(e))
            return JsonResponse(
                {"success": False, "info": "Kindly try again --p2prx2--"}
            )

    def fetch_lowstock(self):
        try:
            products = Product.objects.values("product_name", "quantity").filter(
                quantity__lt=10
            )
            if products:
                return JsonResponse({"success": True, "data": list(products)})
            else:
                return JsonResponse({"success": False, "info": "No products found"})

        except Exception as e:
            logger.warning(str(e))
            return JsonResponse({"success": False, "info": "An error occurred"})

    def sales_analytics(self):
        try:
            purchases = (
                Purchase.objects.values("product__product_name")
                .annotate(
                    total_purchase_amount=Sum("purchase_price"),
                    total_quantities_sold=Sum("quantity"),
                )
                .order_by("-total_quantities_sold", "-total_purchase_amount")
            )

            if purchases:
                best_selling_products = purchases[
                    :5
                ]  # Get the top 5 best-selling products

                # Additional analytics
                total_revenue = sum(item["total_purchase_amount"] for item in purchases)
                average_order_value = (
                    total_revenue / len(purchases) if len(purchases) > 0 else 0
                )

                return JsonResponse(
                    {
                        "success": True,
                        "best_selling_products": list(best_selling_products),
                        "total_revenue": total_revenue,
                        "average_order_value": f"{average_order_value:.2f}",
                    }
                )
            else:
                return JsonResponse({"success": False, "info": "No purchases found"})

        except Exception as e:
            logger.exception("An error occurred: %s", str(e))
            return JsonResponse({"success": False, "info": "An error occurred"})

    # TODO: add an endpoint to get invoices whenever a purchase is made

    def generate_invoice(self, purchase_uid, cashier):
        try:
            purchase = Purchase.objects.select_related("product").get(
                purchase_uid=purchase_uid
            )
            if purchase:
                product_name = purchase.product.product_name
                quantity = purchase.quantity
                purchase_price = purchase.purchase_price
                cashier = cashier
                document = Document()
                document.add_heading("Payment invoice", 0)

                records = [
                    ("Qty", "Product", "Price", "Total"),  # Table header
                    (
                        quantity,
                        product_name,
                        purchase_price / quantity,
                        purchase_price,
                    ),  # Purchase details
                ]

                table = document.add_table(rows=len(records), cols=4)
                for i, row in enumerate(records):
                    row_cells = table.rows[i].cells
                    for j, value in enumerate(row):
                        row_cells[j].text = str(value)

                document.add_paragraph("Signature: _____________________")

                document.add_paragraph(f"Cashier: {cashier}")

                document.add_page_break()

                document.save("purchase.docx")

        except Exception as e:
            logger.warning(str(e))
            return JsonResponse({"success": False, "info": "Kindly try again "})
